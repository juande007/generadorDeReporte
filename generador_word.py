# -*- coding: utf-8 -*-
# encoding=utf8
import time
import json
import re
import os
import io
import GeneradorDeReporte
from datetime import datetime
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Length
from docx.shared import Pt


class Reportes:
    def __init__(self, cliente):
        self.cliente = cliente

    def datosReporte(self, li):

        self.lista = li
        contador = 0
        for element in li:
            print(li[contador][0])
            contador = + 1
            print("Entro a generador word")

    def recolectorDatos():

        incidente = 'LatePacket.json'
        f = open(
            "C:\\Users\CORE i3\Documents\GitHub\generadorDeReporte/Incidentes/"+incidente, "r")
        contenido = f.read()
        archivoJson = json.loads(contenido)

        for lineaArchivoConfig in archivoJson["nuevoIncident"]:
            analisis = lineaArchivoConfig["Analisis"]
            descripcion = lineaArchivoConfig["Descripcion"]
            workAround = lineaArchivoConfig["Workaround"]
            rootCause = lineaArchivoConfig["Root Cause"]
            recomendacion = lineaArchivoConfig["Recomendacion"]

            print(descripcion)
            print(analisis)
            print(rootCause)
            print ("Pruebas...")
    def generadorWord(self, li):

        # id_sonda = int(raw_input('Ingrese el ID de la sonda: '))
        print(("El cliente es desde word: '"+str(self.cliente)+"'"))

        now = datetime.utcnow()
        date_time = now.strftime("%d-%m-%Y")
        date_timeFile = now.strftime("%Y%m%d")

        # C:/Users/CORE i3/Documents/GitHub/generadorDeReporte
        file = open('Clientes/'+self.cliente+'/'+self.cliente+'.json', "r")
        cont = file.read()
        archJson = json.loads(cont)

        for lineaArchivo in archJson["nombreDeCliente"]:
            contacto = lineaArchivo["nombre_contacto"]

        print('Descarga el doc 1')
        document = Document()
        style = document.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(9.1)
        font.italic = True

        self.lista = li
        contador = 0
        contadores = 0
        filas = 0
        numerodeSondas = len(self.lista)
        numFilas = numerodeSondas+1

        if self.cliente != 'TigoCo':
            tablas = document.add_table(rows=numFilas, cols=3)
            tablas.style = 'TableGrid'
            for element in self.lista:
                print(li[contadores][0])
                sonda = li[contadores][0]
                issue = li[contadores][1]

                arch = open("Clientes/"+self.cliente +
                            "/clientNx/clientID/"+sonda+"/"+sonda+".json", "r")
                contenidos = arch.read()
                jsonSonda = json.loads(contenidos)

                for lineaArchivoSonda in jsonSonda["descripcion"]:
                    descripcion_Sondas = lineaArchivoSonda["descripcionSonda"]
                print(li[contadores][0] + "Esto es lo que llega al for y su contador es:"+str(filas))

                if filas == 0:
                    tablas.rows[0].cells[0].text = "ID"
                    tablas.rows[0].cells[1].text = "Descripcion"
                    tablas.rows[0].cells[2].text = "Evaluacion de soporte"
                    tablas.rows[1].cells[0].text = sonda
                    tablas.rows[1].cells[1].text = descripcion_Sondas
                    filas += 1
                    print(
                        li[contadores][0]+"Esto es lo que llega al for y su contador deberia ser 1: "+str(filas))
                else:
                    tablas.rows[filas].cells[0].text = sonda
                    tablas.rows[filas].cells[1].text = descripcion_Sondas
                filas += 1
                print(li[contadores][0] +
                      "Esto es lo que llega al for y su contador deberia ser 2: "+str(filas))
                contadores += 1

            document.add_paragraph(
                "--------------------------------------------------------------------")
            document.add_paragraph("1. GENERAL INFO")
            document.add_paragraph("GEMALTO SUPPORT TEAM - TECHNICAL REPORT")
            document.add_paragraph("Customer: " + str(self.cliente))
            document.add_paragraph("Contac Name: "+str(contacto))
            document.add_paragraph("Call id: ")
            document.add_paragraph("Date & Time: " + str(date_time))

            for element in self.lista:
                print(li[contador][0])
                sonda = li[contador][0]
                issue = li[contador][1]
                evidencia = li[contador][2]
                imgEvidencia = li[contador][3]

                incidente = issue
                sondaALeer = sonda
                evidencias = evidencia
                imagen = imgEvidencia
                print(("La evidencia es: "+evidencias))

                f = open(".\Incidentes/"+incidente, "r")
                contenido = f.read()
                archivoJson = json.loads(contenido)

                for lineaArchivoConfig in archivoJson["nuevoIncident"]:
                    analisis = lineaArchivoConfig["Analisis"]
                    descripcion = lineaArchivoConfig["Descripcion"]
                    workAround = lineaArchivoConfig["Workaround"]
                    rootCause = lineaArchivoConfig["Root Cause"]
                    recomendacion = lineaArchivoConfig["Recomendacion"]

                    print(analisis)

                    arch = open("Clientes/"+self.cliente+"/clientNx/clientID/" +
                                sondaALeer+"/"+sondaALeer+".json", "r")
                    contenidos = arch.read()
                    jsonSonda = json.loads(contenidos)

                    for lineaArchivoSonda in jsonSonda["descripcion"]:
                        descripcion_Sonda = lineaArchivoSonda["descripcionSonda"]

                        document.add_paragraph("")
                        document.add_paragraph("")
                        document.add_paragraph("")
                        document.add_paragraph("")
                        document.add_paragraph(
                            "--------------------------------------------------------------------")
                        paragraph = document.add_paragraph('')
                        run = paragraph.add_run('Cliente ID: ' + str(sonda))
                        run.bold = True
                        paragraph = document.add_paragraph('')
                        run = paragraph.add_run(
                            'Descripcion: ' + descripcion_Sonda)
                        run.bold = True

                        document.add_paragraph(
                            "--------------------------------------------------------------------")
                        document.add_paragraph("2. DETAILS OF INCIDENT: ")
                        document.add_paragraph("Impacted platform: NxClient")
                        document.add_paragraph("Root Cause: " + rootCause)
                        document.add_paragraph(
                            "Incident description: " + descripcion)
                        document.add_paragraph("Evidencias: ")
                        # document.add_picture(imagen)
                        tabla = document.add_table(rows=1, cols=1)
                        tabla.style = 'TableGrid'
                        tabla.rows[0].cells[0].text = evidencias
                        # evidencias+imagen
                        document.add_paragraph("")
                        document.add_paragraph("")

                        document.add_paragraph(
                            "--------------------------------------------------------------------")
                        document.add_paragraph("3. RESOLUTION")
                        document.add_paragraph("Incident Analysis: " + analisis)
                        document.add_paragraph("Workaround: " + workAround)
                        document.add_paragraph("Recommendation: " + recomendacion)
                        document.add_paragraph("Additional comments: NA")

                        contador += 1

            ruta = './Clientes/' + self.cliente+'/Reporte_Errores/'
            print(ruta)
            # 20200918_ATT_Probes
            if os.path.exists(ruta):
                document.save(ruta+'/'+str(date_timeFile)+'_ATT_Probes.docx')
            else:
                print("Entro a else")
                os.makedirs(ruta)
                document.save(ruta+'/'+str(date_timeFile)+'_ATT_Probes.docx')
        else:

            for element in self.lista:
                print(li[contador][0])
                sonda = li[contador][0]
                issue = li[contador][1]
                evidencia = li[contador][2]

                document.add_paragraph("--------------------------------------------------------------------")
                document.add_paragraph("")
                document.add_paragraph("TIGOCO_M_P_" + str(sonda))
                document.add_paragraph("Client ID: ")
                document.add_paragraph("Team Viewer ID: ")
                document.add_paragraph("Fecha Revisión")
                document.add_paragraph("Código de falla: CMO")
                document.add_paragraph("")
                document.add_paragraph("Tareas realizadas para mitigación de fallos: ")
                document.add_paragraph("Las pruebas generales de funcionamiento realizadas")
                document.add_paragraph("")
                document.add_paragraph("--------------------------------------------------------------------")

                contador += 1

            ruta = './Clientes/' + self.cliente+'/Reporte_Errores/'
            print(ruta)
            # 20200918_ATT_Probes
            if os.path.exists(ruta):
                document.save(ruta+'/'+str(date_timeFile)+'_ATT_Probes.docx')
            else:
                print("Entro a else")
                os.makedirs(ruta)
                document.save(ruta+'/'+str(date_timeFile)+'_ATT_Probes.docx')


class DatosReporte:
    def __init__(self, sonda, issue):
        self.sonda = sonda
        self.issue = issue
        self.lista = []


def main():
    reportes = Reportes()
    reportes.generadorWord()


    # reportes.datosReporte()
if __name__ == '__main__':
    main()
